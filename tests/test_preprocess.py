from docx import Document
import pytest

from echr.steps.preprocess_documents import para_to_text, json_table_to_text

class TestPreprocessWord:

    @staticmethod
    @pytest.fixture
    def prepare():
        file = 'tests/data/judgments/001-83979.docx'
        expected_file = 'tests/data/judgments/001-83979_without_smarttags.docx'
        doc = Document(file)
        expected_doc = Document(expected_file)
        broken_list_paragraphs = []
        fixed_list_paragraphs = []
        expected_list_paragraphs = []

        for p in doc.paragraphs:
            broken_list_paragraphs.append(p.text)
            fixed_list_paragraphs.append(para_to_text(p))

        for p in expected_doc.paragraphs:
            expected_list_paragraphs.append(para_to_text(p))

        broken_list_paragraphs = [p for p in broken_list_paragraphs if p]
        fixed_list_paragraphs = [p for p in fixed_list_paragraphs if p]
        expected_list_paragraphs = [p for p in expected_list_paragraphs if p]
        return {'broken': broken_list_paragraphs, 'fixed': fixed_list_paragraphs, 'expected': expected_list_paragraphs}

    @staticmethod
    def test_len_of_paragraphs(prepare):
        assert len(prepare['fixed']) == len(prepare['expected'])

    @staticmethod
    def test_fixed_equals_expected(prepare):
        assert all([p == prepare['fixed'][i] for i, p in enumerate(prepare['expected'])])

    @staticmethod
    def test_broken_different_from_fixed(prepare):
        assert any([p != prepare['fixed'][i] for i, p in enumerate(prepare['broken'])])

    @staticmethod
    def test_broken_different_from_expected(prepare):
        assert any([p != prepare['expected'][i] for i, p in enumerate(prepare['broken'])])


class TestProcessTableAttachment:
    @staticmethod
    def test_json_table_to_text_emtpy():
        table = []
        expected = ""
        res = json_table_to_text(table)
        assert res == expected

    @staticmethod
    def test_json_table_to_text_text_only():
        table = [
            {'header1': 'val1', 'header2': 'val2'},
            {'header1': 'val3', 'header2': 'val4'}]
        expected = "header1 header2\nval1 val2\nval3 val4\n"
        res = json_table_to_text(table)
        assert res == expected

    @staticmethod
    def test_json_table_to_text_int_and_float():
        table = [
            {'header1': 'val1', 'header2': 10},
            {'header1': 12., 'header2': 'val4'}]
        expected = "header1 header2\nval1 10\n12.0 val4\n"
        res = json_table_to_text(table)
        assert res == expected